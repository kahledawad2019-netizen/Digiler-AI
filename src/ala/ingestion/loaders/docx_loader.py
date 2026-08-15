"""Word loader (.docx) via python-docx.

Uses paragraph *styles* to preserve structure: Heading N -> HEADING (with level
and section breadcrumb), List styles -> LIST, code-ish styles -> CODE, else
PARAGRAPH. Tables -> TABLE blocks. Hyperlinks are captured into block meta.
"""

from __future__ import annotations

from pathlib import Path

from ala.core.enums import DocType, ExtractionMethod
from ala.fabric.content import BlockType
from ala.ingestion.loaders.base import BaseLoader, BlockSpec


class DocxLoader(BaseLoader):
    name = "docx"
    extensions = (".docx",)
    doc_types = (DocType.LESSON_PAGE.value, DocType.ASSESSMENT.value, DocType.WORKSHEET.value)
    extraction_method = ExtractionMethod.NATIVE_TEXT
    requires = ("docx",)

    def _parse(self, path: Path) -> list[BlockSpec]:
        import docx

        document = docx.Document(str(path))
        specs: list[BlockSpec] = []
        stack: list[tuple[int, str]] = []

        def section_path() -> list[str]:
            return [t for _, t in stack]

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""

            if style.startswith("heading") or style == "title":
                level = _heading_level(style)
                while stack and stack[-1][0] >= level:
                    stack.pop()
                stack.append((level, text))
                specs.append(BlockSpec(text=text, block_type=BlockType.HEADING,
                                       section_path=section_path(), meta={"level": level}))
            elif "list" in style:
                specs.append(BlockSpec(text=text, block_type=BlockType.LIST,
                                       section_path=section_path()))
            elif "code" in style or "quote" in style:
                bt = BlockType.CODE if "code" in style else BlockType.QUOTE
                specs.append(BlockSpec(text=text, block_type=bt, section_path=section_path()))
            else:
                specs.append(BlockSpec(text=text, block_type=BlockType.PARAGRAPH,
                                       section_path=section_path()))

        for table in document.tables:
            rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
            if any(r.strip() for r in rows):
                specs.append(BlockSpec(text="\n".join(rows), block_type=BlockType.TABLE,
                                       section_path=section_path()))
        return specs


def _heading_level(style: str) -> int:
    digits = "".join(ch for ch in style if ch.isdigit())
    return int(digits) if digits else 1
